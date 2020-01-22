# https://python-docx.readthedocs.io/en/latest/#python-docx

from docx import Document

document = Document(docx='')

paragraph = document.add_paragraph('Something great')
paragraph.add_run('bold')
paragraph.add_run('another text').italic = True

document.add_heading('Test', level=1)